from docx import Document
from docx.shared import Pt,Inches,RGBColor
from PIL import Image
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

df1 = pd.DataFrame({
    '1':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '2':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '3':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '4':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '5':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
})



# Create a new Document
doc = Document()

mainheadPara = doc.add_paragraph()

# page 1 head 
mainheadPararun = mainheadPara.add_run('IMPROVEMENT IN ROAD SIGNAGES ON NATIONAL HIGHWAY IN INDIA USING ARTIFICIAL INTELLIGENCE (AI) BASED SURVEYS')

mainheadPararun.font.name = 'Times New Roman'
mainheadPararun.font.size = Pt(18)
mainheadPararun.font.bold = True

mainheadPara.alignment = 1 


absolute_width_inch_comp = 14.89 * 0.393701  # Absolute width in inches
absolute_height_inch_comp = 5.57 * 0.393701  # Absolute height in inches



image_path_comp = "assets\page1pic.png"

image = Image.open(image_path_comp)

original_width_comp, original_height_comp = image.size

# Calculate scaled dimensions
scaled_width_comp = original_width_comp * 0.58  # 58% of original width
scaled_height_comp = original_height_comp * 0.52  # 52% of original height

doc.add_picture(image_path_comp, width=Inches(absolute_width_inch_comp), height=Inches(absolute_height_inch_comp))

roadData = doc.add_paragraph()
# para1.alignment = 3


nameStretch = roadData.add_run('\nName of the stretch -\n')
nameStretch.font.size = Pt(20)
nameStretch.bold = True

nameOfTheStretch = "Six lane road from mohali to patiala via isbt - 43 singhu border"
roadDataRun = roadData.add_run(f'\n{nameOfTheStretch}\n')

roadDataRun.font.size = Pt(18)
roadDataRun.bold = True

projectData = doc.add_paragraph()
# para2.alignment = 3

projectData.add_run(f'\nUnique Project Code (UPC)\t\t-')
UPC = "N/09001/01002/PB "
projectData.add_run(f'{UPC}\n')
projectData.add_run('State \t\t\t\t\t\t\t-')
State = "Punjab"
projectData.add_run(f'{State}\n')
projectData.add_run('Regional Office (RO) \t\t\t\t-')
RO = "Punjab"
projectData.add_run(f'{RO}\n')
projectData.add_run('Project Implementation Unit (PIU) \t-')
PIU = "Amritsar"
projectData.add_run(f'{PIU}\n')

for run in projectData.runs:
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)

doc.add_page_break()


# page3 of the document
doc.add_page_break()

execHead =doc.add_heading('Executive Summary', level=1)

# Modify heading properties
execHeadRun = execHead.runs[0]
execHeadRun.font.name = 'Calibri'
execHeadRun.font.size = Pt(16)
execHeadRun.bold = False  # Make it bold


# Left align the heading
execHead.alignment = 0

execSumPara = doc.add_paragraph()
# execsumm.alignment = 3

execSumRun1 = execSumPara.add_run('National Highways Authority of India (NHAI), under Ministry of Road Transport and Highways (MoRTH), Government of India has signed a Memorandum of Understanding (MoU) with Indraprastha Institute of Information Technology, Delhi (IIITD) to utilize AI based solutions for carrying out Gap Study w.r.t availability and broad condition of road sign on around 25,000 km of National Highways by assessing the difference between the survey findings and the requirement as per the respective CA & latest Codal provisions relevant to high-speed corridors.\n\n')

execSumRun1.font.name = 'Calibri'
execSumRun1.font.size = Pt(11)
signed_On = "30-10-2025"
date_of_commencement = "30-10-2026"

execSumRun2 = execSumPara.add_run(f'The MoU was signed on {signed_On} and project completion period was kept as twelve (12) months. The date of commencement of this project was {date_of_commencement}.\n\n')

execSumRun2.font.name = 'Calibri'
execSumRun2.font.size = Pt(11)

execSumRun31 = execSumPara.add_run('This report summarizes the findings of a gap study conducted on National Highway project ')

execSumRun31.font.name = 'Calibri'
execSumRun31.font.size = Pt(11)

boldExec = "Six Laning of Jalandhar - Amritsar Section of NH-1 from Km 387.100 to Km 407.100 (Bidhipur Dhilwan)."

execSumRun32 = execSumPara.add_run(boldExec)

execSumRun32.font.name = 'Calibri'
execSumRun32.font.size = Pt(11)
execSumRun32.bold = True

execSumRun33data = "The length of project stretch is "
execSumRun33 = execSumPara.add_run(execSumRun33data)

execSumRun33.font.name = 'Calibri'
execSumRun33.font.size = Pt(11)


boldDistance = "20 km"
execSumRun34 = execSumPara.add_run(boldDistance)

execSumRun34.font.name = 'Calibri'
execSumRun34.font.size = Pt(11)
execSumRun34.bold = True

sixLaneLine = "Six Lane with divided carriageway with Service Road"
option = "Flexible/Rigid"
blank ="value"
stageofProject = "DLP/O&M"
execSumRun35 = execSumPara.add_run(f'with “{sixLaneLine}” configuration and the type of pavement is {option}. The construction work of the project was completed on/ date of COD was {blank} . The project is currently under {stageofProject} stage.')

execSumRun35.font.name = 'Calibri'
execSumRun35.font.size = Pt(11)
# execSumRun34.bold = True

date = "3/10/2025"
no_of_roads = "303"
no_of_roads_nhai = "242"
gap_sign_boards = "-61"
gap_signages = "-61"
execSumRun4 = execSumPara.add_run(f'The survey on this project stretch was carried out on {date} by the team of IIIT Delhi. As per survey findings which is based on Artificial Intelligence (AI), there are {no_of_roads} nos. of roads signs including Chevron, Hazard, Cautionary Warning, Prohibitory Mandatory & Informatory Signs, on this stretch. However, as per NHAI record/ approved Road Signage Plan of Contract Agreement, the number of road signages on this project are {no_of_roads_nhai} nos. (data as provided by NHAI). Therefore, a gap of {gap_sign_boards} nos. sign boards have been observed. Accordingly, there is an additional requirement of {gap_signages} nos. of signages as per the existing contract agreement/ NHAI record. \n\n')

execSumRun4.font.name = 'Calibri'
execSumRun4.font.size = Pt(11)

road_signs_required = "382"
road_signs_rsa = "79"

execSumRun5 = execSumPara.add_run('Secondly, this project has also carried out gap study based on the recommendation of certified road safety auditor (RSA). As per RSA recommendation, the number of road signs required on this project are {road_signs_required} nos. Accordingly, there is an additional requirement of {road_signs_rsa} nos. of signages based on the recommendation of Road Safety Auditor (RSA). \n\n')

execSumRun5.font.name = 'Calibri'
execSumRun5.font.size = Pt(11)


execSumRun6 = execSumPara.add_run('The following table presents the summary of Gap study report : \n\n')

execSumRun6.font.name = 'Calibri'
execSumRun6.font.size = Pt(11)

# page3 of the document
doc.add_page_break()

IntroHead =doc.add_heading('1. Introduction', level=1)

# Modify heading properties
IntroHeadRun = IntroHead.runs[0]
IntroHeadRun.font.name = 'Calibri'
IntroHeadRun.font.size = Pt(16)
IntroHeadRun.bold = False  # Make it bold

IntroSubHead1 = doc.add_heading('1.1.Brief description about the MoU', level=2)

IntroSubHeadRun = IntroSubHead1.runs[0]
IntroSubHeadRun.font.name = 'Calibri'
IntroSubHeadRun.font.size = Pt(14)
IntroSubHeadRun.bold = False  # Make it bold

IntroSubHead1para1 = doc.add_paragraph()
# execsumm.alignment = 3

IntroSubHead1para1run1 = IntroSubHead1para1.add_run('National Highways Authority of India (NHAI), under Ministry of Road Transport and Highways (MoRTH), Government of India has signed a Memorandum of Understanding (MoU) with Indraprastha Institute of Information Technology, Delhi (IIITD) to utilize AI based solutions for carrying out Gap Study w.r.t availability and broad condition of road sign on around 25,000 km of National Highways by assessing the difference between the survey findings and the requirement as per the respective CA & latest Codal provisions relevant to high-speed corridors.\n')

IntroSubHead1para1run1.font.name = 'Calibri'
IntroSubHead1para1run1.font.size = Pt(11)

IntroSubHead1para2 = doc.add_paragraph()
IntroSubHead1para2run1 = IntroSubHead1para2.add_run('The project duration as per the MoU is twelve (12) months and the date of commencement of work is 27.09.2024.  The tentative length of road to be covered under the aforementioned study shall be 25,000 km. The list of stretches included in the project are from different states which is divided into 05 zones (Zone A to E).\n\n')

IntroSubHead1para2run1.font.name = 'Calibri'
IntroSubHead1para2run1.font.size = Pt(11)

table = doc.add_table(rows=df1.shape[0] + 1, cols=df1.shape[1])
table.style = 'Table Grid'

# Center align text in all cells
def center_align_cell(cell):
    """Centers text horizontally and vertically in a table cell."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center vertically
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center horizontally


# Add headers to the first row
for j, col_name in enumerate(df1.columns):
    cell = table.cell(0, j)
    table.cell(0, j).text = col_name  # Set header text
    center_align_cell(cell)

# Add DataFrame values row by row
for i, row in df1.iterrows():
    for j, value in enumerate(row):
        table.cell(i + 1, j).text = str(value)  # Insert cell values

# def set_table_borders(table):
#     tbl = table._element
#     tbl_borders = parse_xml(
#         r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
#         r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
#         r'</w:tblBorders>'
#     )
#     tbl.append(tbl_borders)

# set_table_borders(table)

doc.add_page_break()

absolute_width_inch_india = 11.92 * 0.393701  # Absolute width in inches
absolute_height_inch_india = 11.83 * 0.393701  # Absolute height in inches




image_path_india = "assets\indiamap.png"

image_india = Image.open(image_path_india)

original_width_india, original_height_india = image_india.size

# Calculate scaled dimensions
scaled_width_india = original_width_india * 0.60  # 60% of original width
scaled_height_india = original_height_india * 0.53  # 53% of original height

doc.add_picture(image_path_india, width=Inches(absolute_width_inch_india), height=Inches(absolute_height_inch_india))

IntroSubHead2 = doc.add_heading('1.2 Objectives of Gap Study ', level=2)

IntroSubHead2Run = IntroSubHead2.runs[0]
IntroSubHead2Run.font.name = 'Calibri'
IntroSubHead2Run.font.size = Pt(14)
IntroSubHead2Run.bold = False  # Make it bold

IntroSubHead2para1 = doc.add_paragraph()
# execsumm.alignment = 3

IntroSubHead2para1run1 = IntroSubHead2para1.add_run('Adequate availability of the road signs on the roads plays a significant role in the road safety. NHAI intends to enhance the road safety for all road users by embracing innovation and adopting advanced technologies. \n\n')

IntroSubHead2para1run1.font.name = 'Calibri'
IntroSubHead2para1run1.font.size = Pt(11)


IntroSubHead2para1run2 = IntroSubHead2para1.add_run('Artificial Intelligence (Al) has emerged as a powerful tool for automating tasks and improving data analysis capabilities. By harnessing the potential of Al and Geographic Information Systems (GIS), NHAI can revolutionize its approach w.r.t road signs inspection. \n\n')

IntroSubHead2para1run2.font.name = 'Calibri'
IntroSubHead2para1run2.font.size = Pt(11)

IntroSubHead2para1run3 = IntroSubHead2para1.add_run('Therefore, NHAI in collaboration with IIIT Delhi intends to utilize Al based solutions for Improvement in the availability of road signs on National Highways in India. \n\n')

IntroSubHead2para1run3.font.name = 'Calibri'
IntroSubHead2para1run3.font.size = Pt(11)

objPara = doc.add_paragraph()
objPara.alignment = 1

Objectives_from_Mou = "WRITE OBJECTIVES SIGNED FROM MOU"

objParaRun = objPara.add_run(f'{Objectives_from_Mou}')
objParaRun.font.name = 'Calibri'
objParaRun.font.size = Pt(11)


IntroSubHead3 = doc.add_heading('1.3 Scope of services ', level=2)

IntroSubHead3Run = IntroSubHead3.runs[0]
IntroSubHead3Run.font.name = 'Calibri'
IntroSubHead3Run.font.size = Pt(14)
IntroSubHead3Run.bold = False  # Make it bold

IntroSubHead3Para1 = doc.add_paragraph()

IntroSubHead3Para1Run1 = IntroSubHead3Para1.add_run('The scope of work is to carry out Gap studies w.r.t the availability and broad condition of road signs on around 25,000 km of National Highways in India.  The services includes following:  \n\n')



IntroSubHead3Para1Run1.font.name = 'Calibri'
IntroSubHead3Para1Run1.font.size = Pt(11)


IntroSubHead3Para1Run2 = IntroSubHead3Para1.add_run('i. The gap study shall be carried out by assessing the difference between the survey findings and the requirements of road signs as per signage plan of the respective Contract Agreement. \n\n')



IntroSubHead3Para1Run2.font.name = 'Calibri'
IntroSubHead3Para1Run2.font.size = Pt(11)

doc.add_page_break()

IntroSubHead3Para1Run3 = IntroSubHead3Para1.add_run('ii. Gap study based on updated / latest Codal provisions relevant to high-speed corridors to cater for enhanced safety requirements. For this purpose, IIITD shall engage a certified Road Safety Auditor (RSA) for assisting in the gap study and preparation of report.  \n\n')



IntroSubHead3Para1Run3.font.name = 'Calibri'
IntroSubHead3Para1Run3.font.size = Pt(11)

IntroSubHead3Para1Run4 = IntroSubHead3Para1.add_run('iii. IIITD shall carry out the surveys for collecting imagery and other ancillary data related to availability and condition of road signages on selected National Highways stretches in India as provided by NHAI.  \n\n')



IntroSubHead3Para1Run4.font.name = 'Calibri'
IntroSubHead3Para1Run4.font.size = Pt(11)

IntroSubHead3Para1Run5 = IntroSubHead3Para1.add_run('iv. The data collected through surveys shall be processed through deployment of adequately capable Artificial Intelligence (AI) for accurate identification and classification of road sign types.  \n\n')

IntroSubHead3Para1Run5.font.name = 'Calibri'
IntroSubHead3Para1Run5.font.size = Pt(11)

IntroSubHead4 = doc.add_heading('1.4 Salient features of the instant National Highway Project \n\n', level=2)

IntroSubHead3Run = IntroSubHead4.runs[0]
IntroSubHead3Run.font.name = 'Calibri'
IntroSubHead3Run.font.size = Pt(14)
IntroSubHead3Run.bold = False  # Make it bold

IntroSubHead3Para1 = doc.add_paragraph()

IntroSubHead3Para1Run1 = IntroSubHead3Para1.add_run('The salient features of the instant National Highway Project are mentioned as under:\n\n')



IntroSubHead3Para1Run1.font.name = 'Calibri'
IntroSubHead3Para1Run1.font.size = Pt(11)

doc.add_page_break()

index_map_data = doc.add_paragraph()

mapRun1 = index_map_data.add_run('Index map:\n')
mapRun1.font.name = 'Verdana'
mapRun1.font.size = Pt(11)
mapRun1.bold = True

start_chainage = "3"
mapRun2 = index_map_data.add_run(f'Start Chainage\t:{start_chainage}\n')
mapRun2.font.name = 'Verdana'
mapRun2.font.size = Pt(11)
# mapRun2.bold = True

end_chainage = "400"
mapRun3 = index_map_data.add_run(f'End Chainage\t:{end_chainage}\n\n')
mapRun3.font.name = 'Verdana'
mapRun3.font.size = Pt(11)
# maprun1.bold = True

mapRun4 = index_map_data.add_run('Total Service Roads Surveyed -')
mapRun4.font.name = 'Times New Roman'
mapRun4.font.size = Pt(12)
mapRun4.bold = True

Total_Service_Roads_Surveyed = "56"
mapRun4Val = index_map_data.add_run(f'{Total_Service_Roads_Surveyed}\n\n')

mapRun5 = index_map_data.add_run('Total Intersections Surveyed -')
mapRun5.font.name = 'Times New Roman'
mapRun5.font.size = Pt(12)
mapRun5.bold = True

Total_Intersections_Surveyed = "10" 
mapRun5Val = index_map_data.add_run(f'{Total_Intersections_Surveyed}')

#table here 

absolute_width_inch_scrnst1 = 15.6 * 0.393701  # Absolute width in inches
absolute_height_inch_scrnst1 = 10.82 * 0.393701  # Absolute height in inches



image_path_scrnst1 = "assets\scrnst1.png"

image_scrnst1 = Image.open(image_path_scrnst1)

original_width_scrnst1, original_height_scrnst1 = image_scrnst1.size

# Calculate scaled dimensions
scaled_width_scrnst1 = original_width_scrnst1 * 0.75  # 75% of original width
scaled_height_scrnst1 = original_height_scrnst1 * 0.56  # 56% of original height

doc.add_picture(image_path_scrnst1, width=Inches(absolute_width_inch_scrnst1), height=Inches(absolute_height_inch_scrnst1))

doc.add_page_break()

doc.add_page_break()

headMethod =doc.add_heading('2. Methodology', level=1)

# Modify heading properties
runHeadMethod = headMethod.runs[0]
runHeadMethod.font.name = 'Calibri'
runHeadMethod.font.size = Pt(16)
runHeadMethod.bold = False  # Make it bold

mSubHead1 = doc.add_heading('2.1.	Data Collection ', level=2)

mSubHead1Run = mSubHead1.runs[0]
mSubHead1Run.font.name = 'Calibri'
mSubHead1Run.font.size = Pt(13)
mSubHead1Run.bold = False  # Make it bold

mSubHead1Para1 = doc.add_paragraph()
mSubHead1Para1Run1 = mSubHead1Para1.add_run('The research methodology used is a combination of data collection and processing the same through Artificial Intelligence (AI) based solutions. The data collection involves the use of technology like integrated device which is mounted on the vehicle and the survey is done with expert and driver in the vehicle. The start point is marked and the Main Carriageway (MCW) is covered on both LHS & RHS. The survey further includes covering all the service roads and intersections of the same stretch on vehicles to collect data remotely. The data which is collected is uploaded for the further processing through AI models. \n\n')

mSubHead1Para1Run1.font.name = 'Verdana'
mSubHead1Para1Run1.font.size = Pt(11)

mSubHead2 = doc.add_heading('2.2.	Broad Methodology- Artificial Intelligence based solution ', level=2)

mSubHead2 = mSubHead2.runs[0]
mSubHead2.font.name = 'Calibri'
mSubHead2.font.size = Pt(13)
mSubHead2.bold = False  # Make it bold

mSubHead2Para1 = doc.add_paragraph()
mSubHead2Para1Run1 = mSubHead2Para1.add_run('AI models are employed to process the collected data, identify patterns, and generate insights into road signages as seen in the video captured. The chainage, name of the road, latitude & longitude are marked simultaneously.  \n\n')

mSubHead2Para1Run1.font.name = 'Verdana'
mSubHead2Para1Run1.font.size = Pt(11)

doc.add_page_break()

invHead =doc.add_heading('3. Inventory of Road Signs', level=1)

# Modify heading properties
runInvHead = invHead.runs[0]
runInvHead.font.name = 'Calibri'
runInvHead.font.size = Pt(16)
runInvHead.bold = False  # Make it bold

invSubHead1 = doc.add_heading('3.1.Road Signage Inventory (provided by NHAI)', level=2)

invSubHead1Run = invSubHead1.runs[0]
invSubHead1Run.font.name = 'Calibri'
invSubHead1Run.font.size = Pt(13)
invSubHead1Run.bold = False  # Make it bold


invSubHead2 = doc.add_heading('3.2.Road Signage Inventory based on Survey Data (conducted by IIIT)', level=2)

invSubHead2Run = invSubHead2.runs[0]
invSubHead2Run.font.name = 'Calibri'
invSubHead2Run.font.size = Pt(13)
invSubHead2Run.bold = False  # Make it bold

invSubHead3 = doc.add_heading('3.3.Road signage Requirement as per Road Safety Audit', level=2)

invSubHead3Run = invSubHead3.runs[0]
invSubHead3Run.font.name = 'Calibri'
invSubHead3Run.font.size = Pt(13)
invSubHead3Run.bold = False  # Make it bold

doc.add_page_break()


gapHead =doc.add_heading('4. Results of Gap study', level=1)

# Modify heading properties
runGapHead = gapHead.runs[0]
runGapHead.font.name = 'Calibri'
runGapHead.font.size = Pt(16)
runGapHead.bold = False  # Make it bold

gapSubHead1 = doc.add_heading('4.1.Gap study based on NHAI data', level=2)

gapSubHead1Run = gapSubHead1.runs[0]
gapSubHead1Run.font.name = 'Calibri'
gapSubHead1Run.font.size = Pt(13)
gapSubHead1Run.bold = False  # Make it bold


gapSubHead2 = doc.add_heading('4.2.Gap study based on Rad safety audit', level=2)

gapsubHead2Run = gapSubHead2.runs[0]
gapsubHead2Run.font.name = 'Calibri'
gapsubHead2Run.font.size = Pt(13)
gapsubHead2Run.bold = False  # Make it bold

doc.save("aditya.docx")