from plotly.io import write_image
from docx import Document
import io
import plotly.graph_objects as go
from docx.enum.table import WD_ALIGN_VERTICAL
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt,Inches,RGBColor
from PIL import Image
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def remove_margins(doc):
    """ Remove all margins in the Word document to allow images to use the full page """
    section = doc.sections[0]
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0

def save_chart_to_image(fig):
    """ Save Plotly figure as image in memory and return the image file stream """
    img_stream = io.BytesIO()
    img_bytes = fig.to_image(format="png")  # format can be 'png', 'jpeg', 'svg', etc.
    img_stream.write(img_bytes)
    img_stream.seek(0)  # Go back to the beginning of the BytesIO stream
    with open("chart.png", "wb") as f:
        f.write(img_stream.read())
    return img_stream

def create_word_doc(data, file_name="Chainage_Wise_Analyzed_Data.docx"):
    doc = Document()

    doc.add_heading('Chainage Wise Gap Analysis', 0)

    for chainage, (transposed_data1, transposed_data2, gap_analysis, figs) in data.items():
        doc.add_heading(f"Chainage: {chainage}", level=1)
        
        # **Table 1: PIU Data**
        doc.add_heading("PIU Data:", level=2)
        table1 = doc.add_table(rows=1, cols=len(transposed_data1.columns))

        hdr_cells = table1.rows[0].cells
        for i, col in enumerate(transposed_data1.columns):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background_color(hdr_cells[i], 'ADD8E6')

        for index, row in transposed_data1.iterrows():
            row_cells = table1.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
            if index % 2 == 0:
                for cell in row_cells:
                    set_cell_background_color(cell, 'F0F8FF')
            else:
                for cell in row_cells:
                    set_cell_background_color(cell, 'FFFFFF')

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # **Table 2: RA Data**
        doc.add_heading("RA Data:", level=2)
        table2 = doc.add_table(rows=1, cols=len(transposed_data2.columns))

        hdr_cells = table2.rows[0].cells
        for i, col in enumerate(transposed_data2.columns):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background_color(hdr_cells[i], 'ADD8E6')

        for index, row in transposed_data2.iterrows():
            row_cells = table2.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
            if index % 2 == 0:
                for cell in row_cells:
                    set_cell_background_color(cell, 'F0F8FF')
            else:
                for cell in row_cells:
                    set_cell_background_color(cell, 'FFFFFF')

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # **Gap Analysis Table**
        doc.add_heading("Gap Analysis:", level=2)
        gap_table = doc.add_table(rows=1, cols=len(gap_analysis.columns))

        hdr_cells = gap_table.rows[0].cells
        for i, col in enumerate(gap_analysis.columns):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background_color(hdr_cells[i], 'ADD8E6')

        for index, row in gap_analysis.iterrows():
            row_cells = gap_table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
            if index % 2 == 0:
                for cell in row_cells:
                    set_cell_background_color(cell, 'F0F8FF')
            else:
                for cell in row_cells:
                    set_cell_background_color(cell, 'FFFFFF')

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # remove_margins(doc)
        # for fig in figs:
        img_stream = save_chart_to_image(figs)
        doc.add_picture(img_stream, width=Inches(5.5)) 

        doc.add_page_break()
    # Save the document
    doc.save(file_name)

def set_cell_background_color(cell, color_hex):
    """
    This function sets the background color of a table cell using the color hex code.
    The hex code should be in the format 'RRGGBB'.
    """
    # Get the XML element of the cell
    cell_element = cell._element
    
    # Create the shading element
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)  # Set the fill color with the hex code
    cell_element.get_or_add_tcPr().append(shading) 


def create_word_doc_new(Gap_study_report,moretables,morecharts ,data_new, file_name="Chainage_Wise_Analyzed_Data.docx"):

    df1 = pd.DataFrame({
    '1':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '2':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '3':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '4':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    '5':['saukdghfjs','sakdugdb','sakdhjcdsgcv'],
    })
    
    doc = Document()
    ################### Document generation #########
    mainheadPara = doc.add_paragraph()

    # page 1 head 
    mainheadPararun = mainheadPara.add_run('IMPROVEMENT IN ROAD SIGNAGES ON NATIONAL HIGHWAY IN INDIA USING ARTIFICIAL INTELLIGENCE (AI) BASED SURVEYS')

    mainheadPararun.font.name = 'Times New Roman'
    mainheadPararun.font.size = Pt(18)
    mainheadPararun.font.bold = True

    mainheadPara.alignment = 1 


    absolute_width_inch_comp = 14.89 * 0.393701  # Absolute width in inches
    absolute_height_inch_comp = 5.57 * 0.393701  # Absolute height in inches



    image_path_comp = "assets/page1pic.png"

    image = Image.open(image_path_comp)

    original_width_comp, original_height_comp = image.size

    # Calculate scaled dimensions
    scaled_width_comp = original_width_comp * 0.58  # 58% of original width
    scaled_height_comp = original_height_comp * 0.52  # 52% of original height

    doc.add_picture(image_path_comp, width=Inches(absolute_width_inch_comp), height=Inches(absolute_height_inch_comp))

    roadData = doc.add_paragraph()
    # para1.alignment = 3


    nameStretch = roadData.add_run('\nName of the stretch -\n')
    nameStretch.font.size = Pt(16)
    nameStretch.bold = True

    nameOfTheStretch = "Six lane road from mohali to patiala via isbt - 43 singhu border"
    roadDataRun = roadData.add_run(f'\n{nameOfTheStretch}\n')

    roadDataRun.font.size = Pt(16)
    roadDataRun.bold = True

    projectData = doc.add_paragraph()
    # para2.alignment = 3

    projectData.add_run(f'\nUnique Project Code (UPC)\t\t-')
    UPC = "N/09001/01002/PB "
    projectData.add_run(f'{UPC}\n')
    projectData.add_run('State \t\t\t\t\t\t\t-')
    State = "Punjab"
    projectData.add_run(f'{State}\n')
    projectData.add_run('Regional Office (RO) \t\t\t\t-')
    RO = "Punjab"
    projectData.add_run(f'{RO}\n')
    projectData.add_run('Project Implementation Unit (PIU) \t-')
    PIU = "Amritsar"
    projectData.add_run(f'{PIU}\n')

    for run in projectData.runs:
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    doc.add_page_break()

    # page3 of the document
    doc.add_page_break()

    execHead =doc.add_heading('Executive Summary', level=1)

    # Modify heading properties
    execHeadRun = execHead.runs[0]
    execHeadRun.font.name = 'Calibri'
    execHeadRun.font.size = Pt(16)
    execHeadRun.bold = False  # Make it bold


    # Left align the heading
    execHead.alignment = 0

    execSumPara = doc.add_paragraph()
    execSumPara.alignment = 3

    execSumRun1 = execSumPara.add_run('\n\n\tNational Highways Authority of India (NHAI), under Ministry of Road Transport and Highways (MoRTH), Government of India has signed a Memorandum of Understanding (MoU) with Indraprastha Institute of Information Technology, Delhi (IIITD) to utilize AI based solutions for carrying out Gap Study w.r.t availability and broad condition of road sign on around 25,000 km of selected National Highways (provided by NHAI). The gap study shall be carried out by assessing the difference between the AI based survey findings and the road signage inventory as per NHAI record. The other part of the gap study shall  cover the additional requirement of road signages as per the recommendations of certified Road Safety Auditor (RSA) based on latest codal provisions relevant to high-speed corridors. \t\t\t\n\n')

    execSumRun1.font.name = 'Calibri'
    execSumRun1.font.size = Pt(11)
    signed_On = "30-10-2025"
    date_of_commencement = "30-10-2026"

    execSumRun2 = execSumPara.add_run(f'\tThe MoU was signed on <b>{signed_On}</b> and project completion period was kept as twelve (12) months. The date of commencement of this project was {date_of_commencement}.\n\n')

    execSumRun2.font.name = 'Calibri'
    execSumRun2.font.size = Pt(11)

    execSumRun31 = execSumPara.add_run('\tThis report summarizes the findings of a gap study conducted on National Highway project ')

    execSumRun31.font.name = 'Calibri'
    execSumRun31.font.size = Pt(11)

    boldExec = "Six Laning of Jalandhar - Amritsar Section of NH-1 from Km 387.100 to Km 407.100 (Bidhipur Dhilwan)."

    execSumRun32 = execSumPara.add_run(boldExec)

    execSumRun32.font.name = 'Calibri'
    execSumRun32.font.size = Pt(11)
    execSumRun32.bold = True

    execSumRun33data = "The linear length of project stretch is "
    execSumRun33 = execSumPara.add_run(execSumRun33data)

    execSumRun33.font.name = 'Calibri'
    execSumRun33.font.size = Pt(11)


    boldDistance = "20 km"
    execSumRun34 = execSumPara.add_run(boldDistance)

    execSumRun34.font.name = 'Calibri'
    execSumRun34.font.size = Pt(11)
    execSumRun34.bold = True

    execSumRun334data = "and it includes "
    execSumRun334 = execSumPara.add_run(execSumRun334data)

    execSumRun334.font.name = 'Calibri'
    execSumRun334.font.size = Pt(11)

    sixLaneLine = "Six Lane with divided carriageway with Service Road"
    option = "Flexible/Rigid"
    blank ="value"
    stageofProject = "DLP/O&M"
    execSumRun35 = execSumPara.add_run(f'“{sixLaneLine}” configuration and the type of pavement is {option}. The construction work of the project was completed on/ date of COD was {blank} . The project is currently under {stageofProject} stage.')

    execSumRun35.font.name = 'Calibri'
    execSumRun35.font.size = Pt(11)
    # execSumRun34.bold = True

    date = "3/10/2025"
    no_of_roads = "303"
    no_of_roads_nhai = "242"
    gap_sign_boards = "-61"
    gap_signages = "-61"
    execSumRun4 = execSumPara.add_run(f'\tThe survey on this project stretch was carried out on {date} by the team of IIIT Delhi. As per survey findings which is based on Artificial Intelligence (AI), there are {no_of_roads} nos. of roads signs including Chevron, Hazard, Cautionary Warning, Prohibitory Mandatory & Informatory Signs, on this project. However, as per NHAI record/ approved Road Signage Plan of Contract Agreement, the number of road signages on this project are {no_of_roads_nhai} nos. (data as provided by NHAI). Therefore, a gap of {gap_sign_boards} nos. sign boards have been observed vis-à-vis NHAI record. This variance in the count of signages in the contract document and existing signages taken out by AI survey report is not available in NHAI record. This finding will help NHAI in better inventorisation of the existing Road Signage on this project.  \n\n')

    execSumRun4.font.name = 'Calibri'
    execSumRun4.font.size = Pt(11)

    road_signs_required = "382"
    road_signs_rsa = "303"
    road_signs_rsa_gap = "79"


    execSumRun5 = execSumPara.add_run('\tSecondly, this project has also carried out gap study based on the recommendation of certified road safety auditor (RSA). As per RSA recommendation, the number of Total road signs required on this project are {road_signs_required} nos. Accordingly, there is an additional requirement of {road_signs_rsa_gap} of signages over and above {road_signs_rsa} nos. as observed during the survey. \n\n')

    execSumRun5.font.name = 'Calibri'
    execSumRun5.font.size = Pt(11)


    execSumRun6 = execSumPara.add_run('The following table presents the summary of Gap study report : \n\n')

    execSumRun6.font.name = 'Calibri'
    execSumRun6.font.size = Pt(11)

    gap_table = doc.add_table(rows=1, cols=len(Gap_study_report.columns))

    hdr_cells = gap_table.rows[0].cells
    for i, col in enumerate(Gap_study_report.columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[i], 'fcfcfc')

    for index, row in Gap_study_report.iterrows():
        row_cells = gap_table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and val.is_integer():
                row_cells[i].text = str(int(val))
            else:
                row_cells[i].text = str(val)
        if index % 2 == 0:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')
        else:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    ######## 1. Introduction ##############################################################################################################
    IntroHead =doc.add_heading('1. Introduction', level=1)

    # Modify heading properties
    IntroHeadRun = IntroHead.runs[0]
    IntroHeadRun.font.name = 'Calibri'
    IntroHeadRun.font.size = Pt(16)
    IntroHeadRun.bold = False  # Make it bold

    ### Intro Sub Headings 1 ##############################################################################################################

    IntroSubHead1 = doc.add_heading('1.1.Brief description about the MoU', level=2)

    IntroSubHeadRun = IntroSubHead1.runs[0]
    IntroSubHeadRun.font.name = 'Calibri'
    IntroSubHeadRun.font.size = Pt(14)
    IntroSubHeadRun.bold = False  # Make it bold

    IntroSubHead1para1 = doc.add_paragraph()
    IntroSubHead1para1.alignment = 3

    IntroSubHead1para1run1 = IntroSubHead1para1.add_run('\tA Memorandum of Understanding (MoU) has been signed between the National Highways Authority of India (NHAI) and Indraprastha Institute of Information Technology, Delhi (IIIT Delhi) for carrying out Gap Study w.r.t availability and broad condition of road sign on around 25,000 km National Highways in India. \t\t\t\n')

    IntroSubHead1para1run1.font.name = 'Calibri'
    IntroSubHead1para1run1.font.size = Pt(11)

    IntroSubHead1para2 = doc.add_paragraph()
    IntroSubHead1para2run1 = IntroSubHead1para2.add_run('\tThe project duration as per the MoU is twelve (12) months and the date of commencement of work is 27.09.2024.  The tentative length of roads to be covered under the aforementioned study shall be 25,000 km. The list of stretches included in the project are from different states which is divided into 05 zones (Zone A to E).\t\t\t\t')

    IntroSubHead1para2run1.font.name = 'Calibri'
    IntroSubHead1para2run1.font.size = Pt(11)

    zone_table = pd.DataFrame({
        'Zone' :  ['Zone A','Zone B','Zone C','Zone D','Zone E'],
        'States' : ['Jammu & Kashmir, Punjab, Himachal Pradesh, Uttarakhand , Haryana , Chandigarh and Delhi','Gujrat and Rajasthan','Maharashtra and Madhya Pradesh','Andhra Pradesh , Tamil Nadu , Telangana , Karnataka , Kerala , Chhattisgarh','Assam , Bihar , Meghalaya , Jharkhand , Odisha , Uttar Pradesh and West Bengal'],
    })

    zone_table_word = doc.add_table(rows=1, cols=len(zone_table.columns))

    hdr_cells = zone_table_word.rows[0].cells
    for i, col in enumerate(zone_table.columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[i], 'fcfcfc')

    for index, row in zone_table.iterrows():
        row_cells = zone_table_word.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and val.is_integer():
                row_cells[i].text = str(int(val))
            else:
                row_cells[i].text = str(val)
        if index % 2 == 0:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')
        else:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    absolute_width_inch_india = 11.92 * 0.393701  # Absolute width in inches
    absolute_height_inch_india = 11.83 * 0.393701  # Absolute height in inches

    image_path_india = "assets/indiamap.png"

    image_india = Image.open(image_path_india)

    original_width_india, original_height_india = image_india.size

    # Calculate scaled dimensions
    scaled_width_india = original_width_india * 0.60  # 60% of original width
    scaled_height_india = original_height_india * 0.53  # 53% of original height

    doc.add_picture(image_path_india, width=Inches(absolute_width_inch_india), height=Inches(absolute_height_inch_india))

    ### Intro Sub Headings 2 ################################################################################################################
    
    IntroSubHead2 = doc.add_heading('1.2 Objectives of Gap Study ', level=2)

    IntroSubHead2Run = IntroSubHead2.runs[0]
    IntroSubHead2Run.font.name = 'Calibri'
    IntroSubHead2Run.font.size = Pt(14)
    IntroSubHead2Run.bold = False  # Make it bold

    IntroSubHead2para1 = doc.add_paragraph()
    IntroSubHead2para1.alignment = 3

    IntroSubHead2para1run1 = IntroSubHead2para1.add_run('\n\tAdequate availability of the road signs on the roads plays a significant role in the road safety. NHAI intends to enhance the road safety for all road users by embracing innovation and adopting advanced technologies. \t\t\t\n\n')

    IntroSubHead2para1run1.font.name = 'Calibri'
    IntroSubHead2para1run1.font.size = Pt(11)


    IntroSubHead2para1run2 = IntroSubHead2para1.add_run('\tArtificial Intelligence (Al) has emerged as a powerful tool for automating tasks and improving data analysis capabilities. By harnessing the potential of Al and Geographic Information Systems (GIS), NHAI can revolutionize its approach w.r.t road signs inspection. \t\t\t\n\n')

    IntroSubHead2para1run2.font.name = 'Calibri'
    IntroSubHead2para1run2.font.size = Pt(11)

    IntroSubHead2para1run3 = IntroSubHead2para1.add_run('\tTherefore, NHAI in collaboration with IIIT Delhi intends to utilize Al based solutions for Improvement in the availability of road signs on National Highways in India. \t\t\t\n\n')

    IntroSubHead2para1run3.font.name = 'Calibri'
    IntroSubHead2para1run3.font.size = Pt(11)

    # objPara = doc.add_paragraph()
    # objPara.alignment = 1

    # Objectives_from_Mou = "WRITE OBJECTIVES SIGNED FROM MOU"

    # objParaRun = objPara.add_run(f'{Objectives_from_Mou}')
    # objParaRun.font.name = 'Calibri'
    # objParaRun.font.size = Pt(11)

    ### Intro Sub Headings 3 ################################################################################################################
    
    IntroSubHead3 = doc.add_heading('1.3 Scope of services ', level=2)

    IntroSubHead3Run = IntroSubHead3.runs[0]
    IntroSubHead3Run.font.name = 'Calibri'
    IntroSubHead3Run.font.size = Pt(14)
    IntroSubHead3Run.bold = False  # Make it bold

    IntroSubHead3Para1 = doc.add_paragraph()
    IntroSubHead3Para1.alignment = 3


    IntroSubHead3Para1Run1 = IntroSubHead3Para1.add_run('The scope of work is to carry out Gap studies w.r.t the availability and broad condition of road signs on around 25,000 km of National Highways in India.  The services includes following: \t\t\t\t\n')
    IntroSubHead3Para1Run1.font.name = 'Calibri'
    IntroSubHead3Para1Run1.font.size = Pt(11)


    IntroSubHead3Para1Run2 = IntroSubHead3Para1.add_run('i. The gap study shall be carried out by assessing the difference between the survey findings and the requirements of road signs as per signage plan of the respective Contract Agreement. \n\n')
    IntroSubHead3Para1Run2.font.name = 'Calibri'
    IntroSubHead3Para1Run2.font.size = Pt(11)

    doc.add_page_break()

    IntroSubHead3Para1Run3 = IntroSubHead3Para1.add_run('ii. Gap study based on updated / latest Codal provisions relevant to high-speed corridors to cater for enhanced safety requirements. For this purpose, IIITD shall engage a certified Road Safety Auditor (RSA) for assisting in the gap study and preparation of report.  \n\n')
    IntroSubHead3Para1Run3.font.name = 'Calibri'
    IntroSubHead3Para1Run3.font.size = Pt(11)

    IntroSubHead3Para1Run4 = IntroSubHead3Para1.add_run('iii. IIITD shall carry out the surveys for collecting imagery and other ancillary data related to availability and condition of road signages on selected National Highways stretches in India as provided by NHAI.  \n\n')
    IntroSubHead3Para1Run4.font.name = 'Calibri'
    IntroSubHead3Para1Run4.font.size = Pt(11)

    IntroSubHead3Para1Run5 = IntroSubHead3Para1.add_run('iv. The data collected through surveys shall be processed through deployment of adequately capable Artificial Intelligence (AI) for accurate identification and classification of road sign types.  \n\n')
    IntroSubHead3Para1Run5.font.name = 'Calibri'
    IntroSubHead3Para1Run5.font.size = Pt(11)

    ### Intro Sub Headings 4 ################################################################################################################
    salientFeatures = pd.DataFrame({
        'Description/ Particular':['Name of the project Highway','Unique Project Code (UPC)','NHAI Regional Office','NHAI Project Implementation Unit (PIU)','Details of Project Construction; \n(a) Date of Completion/ COD\n(b) DLP Completion Date','Details of O&M Agency (if any);\n(a) O&M Contract Duration'],
        'Details':['','','','','',''],
    })
    keyfeatures = pd.DataFrame({
        'Description/ Particular': ['Pavement Type &amp; Lane Configuration','Length of Main Carriageway','Length of Service/ Slip Roads','Major Junctions','Minor Junctions','Bridges (Major/Minor)','ROB/RUB/Flyovers','Underpasses (VUP/PUP/CUP)','Culverts','Toll Plaza'],
        'Details': ['','','','','','','','','',''],
    })
    
    IntroSubHead4 = doc.add_heading('1.4 Salient features of the instant National Highway Project \n\n', level=2)

    IntroSubHead4Run = IntroSubHead4.runs[0]
    IntroSubHead4Run.font.name = 'Calibri'
    IntroSubHead4Run.font.size = Pt(14)
    IntroSubHead4Run.bold = False  # Make it bold

    IntroSubHead4Para1 = doc.add_paragraph()
    # IntroSubHead3para1.alignment = 3
    IntroSubHead4Para1Run1 = IntroSubHead4Para1.add_run('i. The salient features of the instant National Highway Project are mentioned as under:\n')
    IntroSubHead4Para1Run1.font.name = 'Calibri'
    IntroSubHead4Para1Run1.font.size = Pt(12)
    IntroSubHead4Para1Run1.bold = True

    salientFeatures_word = doc.add_table(rows=1, cols=len(salientFeatures.columns))

    hdr_cells = salientFeatures_word.rows[0].cells
    for i, col in enumerate(salientFeatures.columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[i], 'fcfcfc')

    for index, row in salientFeatures.iterrows():
        row_cells = salientFeatures_word.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and val.is_integer():
                row_cells[i].text = str(int(val))
            else:
                row_cells[i].text = str(val)
        if index % 2 == 0:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')
        else:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    IntroSubHead4Para2 = doc.add_paragraph()
    # IntroSubHead3para1.alignment = 3
    IntroSubHead4Para1Run2 = IntroSubHead4Para2.add_run('\n\nii. Key Features of the Project:\n')
    IntroSubHead4Para1Run2.font.name = 'Calibri'
    IntroSubHead4Para1Run2.font.size = Pt(12)
    IntroSubHead4Para1Run2.bold = True


    keyfeatures_word = doc.add_table(rows=1, cols=len(keyfeatures.columns))

    hdr_cells = keyfeatures_word.rows[0].cells
    for i, col in enumerate(keyfeatures.columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[i], 'fcfcfc')

    for index, row in keyfeatures.iterrows():
        row_cells = keyfeatures_word.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and val.is_integer():
                row_cells[i].text = str(int(val))
            else:
                row_cells[i].text = str(val)
        if index % 2 == 0:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')
        else:
            for cell in row_cells:
                set_cell_background_color(cell, 'fcfcfc')

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    index_map_data = doc.add_paragraph()
    mapRun1 = index_map_data.add_run('\n\niii. Index map:\n')
    mapRun1.font.name = 'Calibri'
    mapRun1.font.size = Pt(12)
    mapRun1.bold = True

    mapRun2 = index_map_data.add_run('\tIndex map showing the start point and end point of the said project.\n')
    mapRun2.font.name = 'Calibri'
    mapRun2.font.size = Pt(11)
    mapRun2.bold = True

    absolute_width_inch_scrnst1 = 15.6 * 0.393701  # Absolute width in inches
    absolute_height_inch_scrnst1 = 10.82 * 0.393701  # Absolute height in inches
    image_path_scrnst1 = "assets/scrnst1.png"

    image_scrnst1 = Image.open(image_path_scrnst1)

    original_width_scrnst1, original_height_scrnst1 = image_scrnst1.size

    # Calculate scaled dimensions
    scaled_width_scrnst1 = original_width_scrnst1 * 0.75  # 75% of original width
    scaled_height_scrnst1 = original_height_scrnst1 * 0.56  # 56% of original height

    doc.add_picture(image_path_scrnst1, width=Inches(absolute_width_inch_scrnst1), height=Inches(absolute_height_inch_scrnst1))

    IntroSubHead4Para4 = doc.add_paragraph()
    # IntroSubHead3para1.alignment = 3
    IntroSubHead4Para1Run4 = IntroSubHead4Para4.add_run('\n\niv. Total Road segments in this project: \n')
    IntroSubHead4Para1Run4.font.name = 'Calibri'
    IntroSubHead4Para1Run4.font.size = Pt(12)
    IntroSubHead4Para1Run4.bold = True

    line1 = IntroSubHead4Para4.add_run('Main Carraige Way LHS Surveyed -')
    line1.font.name = 'Calibri'
    line1.font.size = Pt(11)
    line1.bold = True

    Total_Service_Roads_Surveyed = "56"
    line1Val = IntroSubHead4Para4.add_run(f'{Total_Service_Roads_Surveyed}\n\n')

    line2 = IntroSubHead4Para4.add_run('Main Carraige Way RHS Surveyed -')
    line2.font.name = 'Calibri'
    line2.font.size = Pt(11)
    line2.bold = True

    Total_Service_Roads_Surveyed = "56"
    line24Val = IntroSubHead4Para4.add_run(f'{Total_Service_Roads_Surveyed}\n\n')
    
    line3 = IntroSubHead4Para4.add_run('Total Service Roads Surveyed -')
    line3.font.name = 'Calibri'
    line3.font.size = Pt(11)
    line3.bold = True

    Total_Service_Roads_Surveyed = "56"
    line3Val = IntroSubHead4Para4.add_run(f'{Total_Service_Roads_Surveyed}\n\n')

    line4 = IntroSubHead4Para4.add_run('Total Intersections Surveyed -')
    line4.font.name = 'Calibri'
    line4.font.size = Pt(11)
    line4.bold = True

    Total_Intersections_Surveyed = "10" 
    line4Val = IntroSubHead4Para4.add_run(f'{Total_Intersections_Surveyed}')


    ### A table here Total Road segments in this project

    doc.add_page_break()

    ### 2. Methodology ################################################################################################################
    
    mSubHead1 = doc.add_heading('2.1.	Data Collection ', level=2)

    mSubHead1Run = mSubHead1.runs[0]
    mSubHead1Run.font.name = 'Calibri'
    mSubHead1Run.font.size = Pt(13)
    mSubHead1Run.bold = False  # Make it bold

    mSubHead1Para1 = doc.add_paragraph()
    mSubHead1Para1Run1 = mSubHead1Para1.add_run('The research methodology used is a combination of data collection and processing the same through Artificial Intelligence (AI) based solutions. The data collection involves the use of technology like integrated device which is mounted on the vehicle and the survey is done with expert and driver in the vehicle. The start point is marked and the Main Carriageway (MCW) is covered on both LHS & RHS. The survey further includes covering all the service roads and intersections of the same stretch on vehicles to collect data remotely. The data which is collected is uploaded for the further processing through AI models. \n\n')

    mSubHead1Para1Run1.font.name = 'Verdana'
    mSubHead1Para1Run1.font.size = Pt(11)

    mSubHead2 = doc.add_heading('2.2.	Broad Methodology- Artificial Intelligence based solution ', level=2)

    mSubHead2 = mSubHead2.runs[0]
    mSubHead2.font.name = 'Calibri'
    mSubHead2.font.size = Pt(13)
    mSubHead2.bold = False  # Make it bold

    mSubHead2Para1 = doc.add_paragraph()
    mSubHead2Para1Run1 = mSubHead2Para1.add_run('AI models are employed to process the collected data, identify patterns, and generate insights into road signages as seen in the video captured. The chainage, name of the road, latitude & longitude are marked simultaneously.  \n\n')

    mSubHead2Para1Run1.font.name = 'Verdana'
    mSubHead2Para1Run1.font.size = Pt(11)

    doc.add_page_break()

    ### 3. Inventory of Road signs ################################################################################################################
    
    invHead =doc.add_heading('3. Inventory of Road Signs', level=1)

    # Modify heading properties
    runInvHead = invHead.runs[0]
    runInvHead.font.name = 'Calibri'
    runInvHead.font.size = Pt(16)
    runInvHead.bold = False  # Make it bold

    invSubHead1 = doc.add_heading('3.1.Road Signage Inventory (provided by NHAI)', level=2)

    invSubHead1Run = invSubHead1.runs[0]
    invSubHead1Run.font.name = 'Calibri'
    invSubHead1Run.font.size = Pt(13)
    invSubHead1Run.bold = False  # Make it bold


    invSubHead2 = doc.add_heading('3.2.Road Signage Inventory based on Survey Data (conducted by IIIT)', level=2)

    invSubHead2Run = invSubHead2.runs[0]
    invSubHead2Run.font.name = 'Calibri'
    invSubHead2Run.font.size = Pt(13)
    invSubHead2Run.bold = False  # Make it bold

    invSubHead3 = doc.add_heading('3.3.Road signage Requirement as per Road Safety Audit', level=2)

    invSubHead3Run = invSubHead3.runs[0]
    invSubHead3Run.font.name = 'Calibri'
    invSubHead3Run.font.size = Pt(13)
    invSubHead3Run.bold = False  # Make it bold

    doc.add_page_break()

    #### 4. Results of Gap Study ################################################################################################################
    gapHead =doc.add_heading('4. Results of Gap study', level=1)

    # Modify heading properties
    runGapHead = gapHead.runs[0]
    runGapHead.font.name = 'Calibri'
    runGapHead.font.size = Pt(16)
    runGapHead.bold = False  # Make it bold

    gapSubHead1 = doc.add_heading('4.1.Gap study based on NHAI data', level=2)

    gapSubHead1Run = gapSubHead1.runs[0]
    gapSubHead1Run.font.name = 'Calibri'
    gapSubHead1Run.font.size = Pt(13)
    gapSubHead1Run.bold = False  # Make it bold


    gapSubHead2 = doc.add_heading('4.2.Gap study based on Rad safety audit', level=2)

    gapsubHead2Run = gapSubHead2.runs[0]
    gapsubHead2Run.font.name = 'Calibri'
    gapsubHead2Run.font.size = Pt(13)
    gapsubHead2Run.bold = False  # Make it bold

    doc.add_page_break()

    ############################################################################################################
    print('more tables',moretables)
    for key,data in moretables.items():
        doc.add_heading(key, level=2)
        gap_table = doc.add_table(rows=1, cols=len(data.columns))

        hdr_cells = gap_table.rows[0].cells
        for i, col in enumerate(data.columns):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background_color(hdr_cells[i], 'fcfcfc')

        for index, row in data.iterrows():
            row_cells = gap_table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float) and val.is_integer():
                    row_cells[i].text = str(int(val))
                else:
                    row_cells[i].text = str(val)
            if index % 2 == 0:
                for cell in row_cells:
                    set_cell_background_color(cell, 'fcfcfc')
            else:
                for cell in row_cells:
                    set_cell_background_color(cell, 'fcfcfc')

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_page_break()

    for fig in morecharts:
        img_stream = save_chart_to_image(fig)
        doc.add_picture(img_stream, width=Inches(5.5)) 

    doc.add_page_break()

    doc.add_heading('Chainage Wise Gap Analysis', 0)

    for chainage, (data, figs) in data_new.items():
        doc.add_heading(f"Chainage: {chainage}", level=1)
        
        # # **Table 1: PIU Data**
        # doc.add_heading("PIU Data:", level=2)
        # table1 = doc.add_table(rows=1, cols=len(transposed_data1.columns))

        # hdr_cells = table1.rows[0].cells
        # for i, col in enumerate(transposed_data1.columns):
        #     hdr_cells[i].text = col
        #     hdr_cells[i].paragraphs[0].runs[0].bold = True
        #     set_cell_background_color(hdr_cells[i], 'ADD8E6')

        # for index, row in transposed_data1.iterrows():
        #     row_cells = table1.add_row().cells
        #     for i, val in enumerate(row):
        #         row_cells[i].text = str(val)
        #     if index % 2 == 0:
        #         for cell in row_cells:
        #             set_cell_background_color(cell, 'F0F8FF')
        #     else:
        #         for cell in row_cells:
        #             set_cell_background_color(cell, 'FFFFFF')

        #     for cell in row_cells:
        #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # # **Table 2: RA Data**
        # doc.add_heading("RA Data:", level=2)
        # table2 = doc.add_table(rows=1, cols=len(transposed_data2.columns))

        # hdr_cells = table2.rows[0].cells
        # for i, col in enumerate(transposed_data2.columns):
        #     hdr_cells[i].text = col
        #     hdr_cells[i].paragraphs[0].runs[0].bold = True
        #     set_cell_background_color(hdr_cells[i], 'ADD8E6')

        # for index, row in transposed_data2.iterrows():
        #     row_cells = table2.add_row().cells
        #     for i, val in enumerate(row):
        #         row_cells[i].text = str(val)
        #     if index % 2 == 0:
        #         for cell in row_cells:
        #             set_cell_background_color(cell, 'F0F8FF')
        #     else:
        #         for cell in row_cells:
        #             set_cell_background_color(cell, 'FFFFFF')

        #     for cell in row_cells:
        #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # **Gap Analysis Table**
        doc.add_heading("Gap Analysis:", level=2)
        gap_table = doc.add_table(rows=1, cols=len(data.columns))

        hdr_cells = gap_table.rows[0].cells
        for i, col in enumerate(data.columns):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background_color(hdr_cells[i], 'fcfcfc')

        for index, row in data.iterrows():
            row_cells = gap_table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float) and val.is_integer():
                    row_cells[i].text = str(int(val))
                else:
                    row_cells[i].text = str(val)
            if index % 2 == 0:
                for cell in row_cells:
                    set_cell_background_color(cell, 'fcfcfc')
            else:
                for cell in row_cells:
                    set_cell_background_color(cell, 'fcfcfc')

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # remove_margins(doc)
        # for fig in figs:
        img_stream = save_chart_to_image(figs)
        doc.add_picture(img_stream, width=Inches(5.5)) 

        doc.add_page_break()
    # Save the document
    doc.save(file_name)